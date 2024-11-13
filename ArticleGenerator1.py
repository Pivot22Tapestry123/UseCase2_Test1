import traceback
import requests
import streamlit as st
import os
import json
import warnings
import io
from datetime import datetime, timedelta
from crewai import Agent, Task, Crew, Process
from docx import Document
from docx.shared import Pt, RGBColor
from docx import Document as DocxDocument  # Import for reading Word documents

# Suppress warnings
warnings.filterwarnings('ignore')

# Helper function to load and save configurations
def load_config():
    try:
        with open("agent_task_config.json", "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def save_config(config):
    with open("agent_task_config.json", "w") as f:
        json.dump(config, f)

# Function to read content from a Word document
def read_docx(file):
    doc = DocxDocument(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Load persisted configurations at startup
config = load_config()

# Streamlit UI
st.title("Research Article Generator")

# Apply custom CSS for background color and logo position
st.markdown(
    """
    <style>
    body {
        background-color: black;
        color: white;  /* Set text color to white for better visibility on black background */
    }
    .css-18e3th9 {
        padding-top: 0rem;
    }
    .logo {
        position: absolute;
        top: 10px;
        left: 10px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Display the company logo in the top left corner
st.image("assets/logo.png", width=100, output_format="PNG")  # Adjust path as needed

# Concise instructions for naming files
st.write(
    "### File Naming Instructions\n"
    "Please name your files starting with the date in `YYYY-MM-DD` format, followed by an underscore (`_`).\n"
    "Example: `2023-06-15_transcript.docx`"
)

# File uploader to accept both .txt and .docx files
uploaded_files = st.file_uploader("Upload one or more transcript files (TXT or Word)", type=["txt", "docx"], accept_multiple_files=True)

# Display file names
if uploaded_files:
    file_names = [file.name for file in uploaded_files]
    st.write("Uploaded Files:", ", ".join(file_names))

# Checkbox to include all documents by automatically setting the date range
include_all_docs = st.checkbox("Include all documents")

# Date range selectors
if include_all_docs:
    # Automatically set date range to 100 years before today to today
    start_date = (datetime.now() - timedelta(days=365 * 100)).date()  # Convert to date
    end_date = datetime.now().date()  # Convert to date
    st.write(f"Automatically including documents from {start_date} to {end_date}")
else:
    # Allow manual date selection without calling .date() since it's already a date object
    start_date = st.date_input("Select Start Date")
    end_date = st.date_input("Select End Date")

# API Key input
openai_api_key = st.text_input("Enter your OpenAI API Key", type="password")

# Temperature slider
temperature = st.slider("Set the temperature for the output (0 = deterministic, 1 = creative)", min_value=0.0, max_value=1.0, value=0.7)

# Toggle for selecting Standard or Custom Prompts
prompt_mode = st.radio("Select Prompt Mode", ["Standard Prompts", "Custom Prompts"])

# Define prompts for agents and tasks
if 'prompts' not in st.session_state:
    st.session_state['prompts'] = config or {
        "planner": {
            "role": "Content Planner", 
            "goal": "Plan engaging and factually accurate content on the given topic",
            "backstory": (
                "You're responsible for analyzing the transcripts to extract key themes, challenges, "
                "and opportunities discussed by industry leaders. Categorize the insights into major "
                "sections, such as Industry Trends, Technological Impacts, Regulatory Considerations, and Future Outlook. "
                "Use participant quotes strategically to add credibility and depth, ensuring you include specific examples "
                "from relevant companies where applicable. Only use the data mentioned in the eligible files "
                "for report generation and do not incorporate any outside knowledge. "
                "Ensure the report reads naturally and has the polished "
                "feel of a human-written document, with varied sentence structures, a professional tone, and engaging, nuanced language."
            )
        },
        "writer": {
            "role": "Content Writer",
            "goal": "Write insightful and factually accurate research report",
            "backstory": (
                "Your task is to write a comprehensive and engaging research article based on the content "
                "plan provided by the Content Planner. Integrate specific quotes from participants to support "
                "key arguments and provide a balanced view of the opportunities and challenges discussed. "
                "Use evidence-based analysis and maintain a formal yet engaging tone. Structure the content "
                "thematically, addressing each major point with supporting data, expert opinions, and specific "
                "examples. Highlight knowledge gaps and propose strategies for addressing them, ensuring the content "
                "is actionable. Write in a way that feels human and natural, as though crafted by a seasoned technical "
                "writer. Avoid robotic language and ensure the narrative is engaging, relatable, and enriched with "
                "cross-references that connect different sections of the report for a cohesive flow. "
                "End the article with a final 'Conclusion' section, which summarizes key insights without adding further suggestions or recommendations. "
                "Only use the data mentioned in the eligible files for report generation and do not incorporate any outside knowledge."
            )
        },
        "editor": {
            "role": "Editor",
            "goal": "Edit a given blog post",
            "backstory": (
                "Your role is to refine the research article drafted by the Content Writer. Ensure the content "
                "follows journalistic best practices, maintains a formal and professional tone, and is well-structured. "
                "Check for balanced viewpoints and make sure that participant quotes are used effectively. Avoid "
                "controversial statements unless necessary, and ensure the report addresses both benefits and risks. "
                "Focus on coherence, readability, and the logical flow of ideas. Make sure there is no content or "
                "additional sections following the Conclusion. The Conclusion should be the final part of the report, "
                "summarizing key insights without adding any further recommendations or suggestions. "
                "Only use the data mentioned in the eligible files for report generation and do not incorporate any outside knowledge."
            )
        },
        "tasks": {
            "plan": (
                "Analyze the transcripts to extract major themes and plan the content structure. Identify key challenges, "
                "opportunities, and knowledge gaps, and suggest where to include participant quotes. Recommend specific case studies, "
                "examples, or statistics that would enrich the report."
            ),
            "write": (
                "Write a research article based on the content plan, integrating participant quotes, evidence-based analysis, specific examples, "
                "and a balanced discussion of opportunities and risks. Ensure the content is engaging, relatable, and structured to connect different themes. "
                "End the article with a final 'Conclusion' section, which summarizes the report without adding further suggestions or recommendations."
            ),
            "edit": (
                "Review and edit the research article to ensure coherence, proper use of quotes, balanced viewpoints, and adherence to journalistic standards. "
                "Make sure that cross-references are present and that the article ends with a Conclusion section only, with no additional recommendations or suggestions afterward."
            )
        }
    }

# User inputs for each prompt based on mode
st.header("Agent Prompts")

is_editable = prompt_mode == "Custom Prompts"

for agent, prompts in st.session_state['prompts'].items():
    if agent != "tasks":
        st.subheader(f"{agent.capitalize()} Agent")
        prompts["role"] = st.text_input(f"{agent.capitalize()} Role", value=prompts["role"], key=f"{agent}_role", disabled=not is_editable)
        prompts["goal"] = st.text_area(f"{agent.capitalize()} Goal", value=prompts["goal"], key=f"{agent}_goal", disabled=not is_editable)
        prompts["backstory"] = st.text_area(f"{agent.capitalize()} Backstory", value=prompts["backstory"], key=f"{agent}_backstory", disabled=not is_editable)

# Task Descriptions UI
st.header("Task Descriptions")
for task, description in st.session_state['prompts']["tasks"].items():
    st.session_state['prompts']["tasks"][task] = st.text_area(f"{task.capitalize()} Task Description", value=description, key=f"{task}_description", disabled=not is_editable)

# Button to save user modifications
if st.button("Save Configuration"):
    save_config(st.session_state['prompts'])
    st.success("Configuration saved successfully!")

# Combine the content of all eligible files
combined_content = ""

# Process files within the selected date range
for i, uploaded_file in enumerate(uploaded_files, start=1):
    file_date_str = uploaded_file.name.split("_")[0]
    try:
        file_date = datetime.strptime(file_date_str, "%Y-%m-%d").date()  # Convert to date
        if start_date <= file_date <= end_date:
            # Read content based on file type
            if uploaded_file.type == "text/plain":
                file_content = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_content = read_docx(uploaded_file)

            # Append file content to combined_content with clear indicators
            combined_content += f"--- Beginning of content from file {i}: {uploaded_file.name} ---\n"
            combined_content += file_content + "\n"
            combined_content += f"--- End of content from file {i}: {uploaded_file.name} ---\n\n"

    except ValueError:
        st.warning(f"The file {uploaded_file.name} does not have a valid date format in the filename. Skipping this file.")

# Button to view combined content
if st.button("View Combined Content"):
    st.text_area("Combined Content Preview", combined_content, height=300)

# Ensure combined content is not empty before proceeding
if combined_content:
    # Define agents and tasks for processing combined content
    planner = Agent(
        role=st.session_state['prompts']['planner']['role'],
        goal=st.session_state['prompts']['planner']['goal'],
        backstory=st.session_state['prompts']['planner']['backstory'],
        allow_delegation=False,
        verbose=True,
        temperature=temperature,
        openai_api_key=openai_api_key  # Pass API key directly
    )

    # Define a single task for content planning based on the combined content
    task = Task(
        description=st.session_state['prompts']['tasks']['plan'],
        agent=planner,
        inputs=[combined_content],
        expected_output="A comprehensive article based on the provided transcripts."
    )

    crew = Crew(agents=[planner], tasks=[task], verbose=True)
    with st.spinner("Processing all uploaded files for a consolidated report..."):
        result = crew.kickoff()

    # Writer agent for cohesive report generation
    writer = Agent(
        role="Content Writer",
        goal="Write a cohesive research article based on organized sections.",
        backstory=st.session_state['prompts']['writer']['backstory'],
        allow_delegation=False,
        verbose=True
    )

    write_task = Task(
        description=st.session_state['prompts']['tasks']['write'],
        agent=writer,
        expected_output="A well-structured and cohesive research article."
    )

    writer_crew = Crew(agents=[writer], tasks=[write_task], verbose=True)
    with st.spinner("Writing the cohesive research article from combined content..."):
        final_report = writer_crew.kickoff()

    # Display the final report
    st.success("Research article generated successfully!")
    st.markdown(final_report)

    # Generate Word document with specified formatting
    doc = Document()
    
    # Set document margins to 1 inch
    doc_sections = doc.sections
    for section in doc_sections:
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Pt(72)  # 1 inch margin

    # Add content to the document
    doc.add_paragraph("Industry Insights Report", style='Heading 1')
    
    for line in final_report.split('\n'):
        clean_line = line.strip('*')  # Remove asterisks from each line
        p = doc.add_paragraph(clean_line)
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(11)
        p.paragraph_format.alignment = 0  # Left align
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1  # Single line spacing

    # Save document to buffer
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    # Download Word document
    st.download_button(
        label="Download Word Document",
        data=word_buffer.getvalue(),
        file_name="research_article.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.warning("No eligible content found in the selected date range.")

st.markdown("---")
st.markdown("Tapestry Networks")
